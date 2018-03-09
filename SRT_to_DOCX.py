from docx import Document
from docx.shared import Inches
import pysrt

in_file = 'motocultor.srt'
doc_name = in_file.split('.')[0]
out_file = in_file.replace('.srt','2.docx')

subs = pysrt.open(in_file, encoding='ansi')

document = Document()

document.add_paragraph(doc_name.upper()+'.')
document.add_paragraph('(Tiempo de duración: )\n')

document.add_picture('monty-truth.jpg', width=Inches(1.25))

for sub in subs:
    print(sub.start)
    print(sub.text)   
    document.add_paragraph(str(sub.start))
    document.add_paragraph(sub.text + '\n')    

document.save(out_file)